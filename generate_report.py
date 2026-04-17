from docx import Document

doc = Document()

doc.add_heading('Doctor Appointment Booking System - Project Report', level=1)

doc.add_paragraph('1. Problem Statement')

doc.add_paragraph(
    'In the existing manual healthcare appointment systems, patients face long waiting periods, appointment conflicts, and lack of transparency. '
    'Doctors and administrators struggle with schedule management, patient tracking, and document coordination. The goal is to remove manual inefficiencies by building a digital Doctor Appointment Booking System using the MERN stack (MongoDB, Express, React, Node.js).'
)

doc.add_paragraph('2. Objectives')
objs = [
    'Implement role-based access control for patients, doctors, and admins.',
    'Enable patients to search doctors by specialization, view profiles, and book appointments online.',
    'Enable doctors to manage availability, view appointments, and update status.',
    'Enable admins to review and approve doctor registrations, manage users, and oversee the system.',
    'Build API endpoints, secure authentication, and persistent data storage using MongoDB.',
    'Deliver a responsive front-end using React with fast workflow and state management.',
    'Ensure reliability with testing, error handling, and deployability on cloud services.'
]
for o in objs:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph('3. Pipeline of Development (System Workflow)')
steps = [
    'Requirement analysis: collect must-have features (user, doctor, admin workflows, appointment lifecycle).',
    'System design: define architecture, data models (User, Doctor, Appointment, Review, Chat), and REST API routes.',
    'Backend development: implement Express server, MongoDB models, authentication, authorization middleware, and controllers for all roles.',
    'Frontend development: build React components, routes, forms, dashboards, and integrate with backend APIs.',
    'Testing: unit tests, integration tests, and manual user acceptance testing for booking flows and approvals.',
    'Deployment: configure hosting (e.g., Heroku, Vercel), connect to production DB, and continuous integration pipeline.'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('4. Expected Results')
res = [
    'A fully operational Doctor Appointment Booking System for multiple user roles.',
    'Improved appointment scheduling efficiency and reduced administrative workload.',
    'Enhanced user experience with short response times and mobile-friendly interface.',
    'Increased transparency for doctors and patients with clear appointment status updates.',
    'Scalable, maintainable codebase ready for future feature extension (notifications, telehealth, analytics).'
]
for r in res:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('5. Project Summary')

summary = (
    'This project is designed to digitize and simplify the clinic appointment process. It supports patients searching and booking appointments, doctors managing schedules, and admins approving doctor registrations. '
    'It emphasizes security, role-based access, and process automation. The implementation has a modular backend, modern React frontend, and follows a systematic SDLC with testing and deployment readiness.'
)

doc.add_paragraph(summary)

doc.save('Doctor_Appointment_Booking_System_Project_Report.docx')
print('created Doctor_Appointment_Booking_System_Project_Report.docx')
